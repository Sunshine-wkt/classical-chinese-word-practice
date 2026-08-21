from pathlib import Path
import sys
from docx import Document

source_path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path('source/reference/yuwoyusuo-notes.docx')
output_path = Path(sys.argv[2]) if len(sys.argv) > 2 else source_path.with_suffix('.txt')

document = Document(source_path)
lines = []

for paragraph in document.paragraphs:
    text = paragraph.text.strip()
    if text:
        lines.append(text)

for table in document.tables:
    for row in table.rows:
        cells = [cell.text.strip().replace('\n', '／') for cell in row.cells]
        if any(cells):
            lines.append('｜'.join(cells))

output_path.write_text('\n'.join(lines) + '\n', encoding='utf-8')
print(f'已擷取 {len(lines)} 行至 {output_path}')
